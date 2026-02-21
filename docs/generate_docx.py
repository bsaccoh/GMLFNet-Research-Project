"""Generate PROJECT_DOCUMENT.docx from scratch using python-docx."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import os


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, header_color="1F4E79"):
    """Add a formatted table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    return table


def create_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("GMLFNet")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(
        "Gradient-Based Meta-Learning with Fast Adaptation Weights\n"
        "for Robust Multi-Centre Polyp Segmentation"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph("")

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph("")

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = type_p.add_run("Comprehensive Project Document")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph("")

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run("Master's Thesis in Computer Science")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph("")

    repo_p = doc.add_paragraph()
    repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = repo_p.add_run("Repository: https://github.com/bsaccoh/GMLFNet-Research-Project")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    toc_heading = doc.add_heading('Table of Contents', level=1)

    toc_items = [
        "1.  Project Purpose",
        "2.  Problem Statement",
        "3.  Scope",
        "4.  Area of Coverage",
        "5.  Model Architecture",
        "6.  Fast Adaptation Weights (FAW) — Novel Contribution",
        "7.  Meta-Learning Framework",
        "8.  Advantages Over Existing Models",
        "9.  Datasets Used",
        "10. Training Pipeline",
        "11. Evaluation Methodology",
        "12. Experimental Design",
        "13. Technical Specifications",
        "14. Project Structure",
        "15. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # SECTION 1: PROJECT PURPOSE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('1. Project Purpose', level=1)

    doc.add_paragraph(
        "GMLFNet is a Master's thesis research project that proposes a novel deep learning "
        "framework for colorectal polyp segmentation that can generalize robustly across "
        "images from multiple medical centres (hospitals) without performance degradation."
    )

    doc.add_paragraph(
        "Colorectal cancer is the third most common cancer worldwide, and early detection "
        "of polyps during colonoscopy is critical for prevention. Automated polyp segmentation "
        "assists gastroenterologists by highlighting polyp regions in real-time during "
        "endoscopic procedures. However, existing segmentation models trained on data from "
        "one hospital often fail when deployed at a different hospital due to domain shift — "
        "differences in imaging equipment, patient demographics, imaging protocols, lighting "
        "conditions, and endoscope types."
    )

    doc.add_paragraph("GMLFNet addresses this challenge by combining:")

    bullets = [
        "A state-of-the-art segmentation architecture (multi-scale encoder-decoder with reverse attention)",
        "A novel Fast Adaptation Weights (FAW) module that enables rapid domain adaptation",
        "MAML-based gradient meta-learning that trains the model to quickly adapt to new domains with minimal data",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        "The ultimate goal is a polyp segmentation model that can be deployed at any new hospital "
        "and adapt within seconds using just a handful of local samples."
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 2: PROBLEM STATEMENT
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('2. Problem Statement', level=1)

    doc.add_heading('2.1 The Domain Shift Problem in Medical Imaging', level=2)

    doc.add_paragraph(
        "Medical imaging data varies significantly across institutions due to multiple factors:"
    )

    add_styled_table(doc,
        ["Factor", "Impact"],
        [
            ["Endoscope manufacturer", "Different sensors produce different color profiles, resolutions, and noise patterns"],
            ["Imaging protocol", "Varying zoom levels, illumination settings, and preparation procedures"],
            ["Patient demographics", "Different population characteristics across regions"],
            ["Image quality", "Differences in sharpness, contrast, and artifacts"],
            ["Annotation style", "Variations in how pathologists delineate polyp boundaries"],
        ]
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "When a deep learning model trained on Hospital A's data is applied to Hospital B's images, "
        "performance typically drops by 10–25% in Dice score. This makes single-centre trained models "
        "unreliable for clinical deployment across healthcare systems."
    )

    doc.add_heading('2.2 Limitations of Existing Approaches', level=2)

    limitations = [
        ("Standard supervised learning", "Trains on pooled multi-centre data but treats all centres equally, failing to capture centre-specific characteristics."),
        ("Domain adaptation methods", "Require access to target domain data during training, which may not be available."),
        ("Fine-tuning", "Adapting the full model on small target datasets leads to catastrophic forgetting and overfitting."),
        ("Existing meta-learning for segmentation", "Adapts all model parameters in the inner loop, which is computationally expensive and prone to overfitting on small support sets."),
    ]
    for title, desc in limitations:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(desc)

    # ════════════════════════════════════════════════════════════════
    # SECTION 3: SCOPE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('3. Scope', level=1)

    doc.add_heading('3.1 In Scope', level=2)
    in_scope = [
        "Polyp segmentation in colonoscopy images from five standardized benchmark datasets representing different medical centres",
        "Meta-learning framework using Model-Agnostic Meta-Learning (MAML) for cross-centre generalization",
        "Novel FAW module as the primary thesis contribution for efficient domain adaptation",
        "Two backbone architectures for comparison: Res2Net-50 (CNN-based) and PVTv2-B2 (Transformer-based)",
        "Comprehensive evaluation including zero-shot generalization, few-shot adaptation, and ablation studies",
        "Reproducible research: full codebase, configuration files, and training notebooks",
    ]
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Out of Scope', level=2)
    out_scope = [
        "Real-time video segmentation (frame-by-frame processing)",
        "3D volumetric segmentation",
        "Other gastrointestinal pathologies (ulcers, bleeding, tumours)",
        "Clinical deployment and regulatory approval",
        "Hardware-specific optimization (TensorRT, ONNX conversion)",
    ]
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')

    # ════════════════════════════════════════════════════════════════
    # SECTION 4: AREA OF COVERAGE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('4. Area of Coverage', level=1)

    doc.add_paragraph(
        "GMLFNet operates at the intersection of several research areas:"
    )

    areas = [
        ("4.1 Computer-Aided Diagnosis (CAD) in Gastroenterology",
         ["Automated detection and delineation of colorectal polyps",
          "Assists endoscopists during colonoscopy procedures",
          "Contributes to early colorectal cancer screening"]),
        ("4.2 Domain Generalization in Medical Imaging",
         ["Training models that transfer across institutional boundaries",
          "Handling distribution shifts in medical data without target domain access",
          "Clinically relevant for deploying AI across hospital networks"]),
        ("4.3 Meta-Learning for Medical Applications",
         ["Few-shot adaptation to new clinical environments",
          "Learning-to-learn paradigm applied to segmentation tasks",
          "Gradient-based meta-learning (MAML family) for rapid adaptation"]),
        ("4.4 Semantic Segmentation",
         ["Pixel-level binary classification (polyp vs. background)",
          "Multi-scale feature extraction and boundary refinement",
          "Deep supervision for training stability"]),
        ("4.5 Feature Modulation and Conditional Computation",
         ["FiLM (Feature-wise Linear Modulation) for domain conditioning",
          "Lightweight parameter generation for task-specific adaptation",
          "Efficient architecture design for resource-constrained environments"]),
    ]
    for heading, bullets in areas:
        doc.add_heading(heading, level=2)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')

    # ════════════════════════════════════════════════════════════════
    # SECTION 5: MODEL ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('5. Model Architecture', level=1)

    doc.add_heading('5.1 Architecture Overview', level=2)

    doc.add_paragraph(
        "The GMLFNet architecture follows a three-stage pipeline:"
    )

    arch_p = doc.add_paragraph()
    arch_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = arch_p.add_run(
        "Input Image (3×352×352)\n"
        "       ↓\n"
        "ENCODER (Res2Net-50 or PVTv2-B2)\n"
        "       ↓\n"
        "Multi-Scale Features [f1, f2, f3, f4]\n"
        "       ↓\n"
        "FAST ADAPTATION WEIGHTS (FAW)\n"
        "   → FiLM modulation (γ, β) per decoder layer\n"
        "       ↓\n"
        "MULTI-SCALE DECODER (RFB + Reverse Attention)\n"
        "       ↓\n"
        "Segmentation Map (1×352×352) + 3 Side Outputs"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('5.2 Encoder (Backbone)', level=2)

    doc.add_paragraph("Two backbone options are provided for comparative analysis:")

    doc.add_heading('Res2Net-50 (CNN-based)', level=3)
    res2net_items = [
        "Multi-scale representation at granular level via Res2Net blocks",
        "Pretrained on ImageNet-1K",
        "Output channels: [256, 512, 1024, 2048] at strides [4, 8, 16, 32]",
        "Total parameters: ~25M",
        "Strengths: Strong local feature extraction, well-established in medical imaging",
    ]
    for item in res2net_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('PVTv2-B2 (Transformer-based)', level=3)
    pvt_items = [
        "Pyramid Vision Transformer with spatial-reduction attention",
        "Pretrained on ImageNet-1K",
        "Output channels: [64, 128, 320, 512] at strides [4, 8, 16, 32]",
        "Total parameters: ~25M",
        "Strengths: Global context modelling, long-range dependencies, state-of-the-art on polyp benchmarks",
    ]
    for item in pvt_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.3 Multi-Scale Decoder', level=2)

    doc.add_paragraph("The decoder incorporates three key components inspired by PraNet:")

    doc.add_heading('Receptive Field Block (RFB)', level=3)
    doc.add_paragraph(
        "Four parallel convolutional branches with dilation rates [1, 3, 5] plus a 1×1 branch. "
        "Captures multi-scale contextual information at each feature level. Applied to f2, f3, f4 "
        "encoder features with residual connection for gradient flow."
    )

    doc.add_heading('Partial Decoder', level=3)
    doc.add_paragraph(
        "Aggregates the three RFB-enhanced features (f2, f3, f4) through upsampling and concatenation. "
        "Produces an initial coarse segmentation prediction with two-layer refinement using 3×3 convolutions."
    )

    doc.add_heading('Reverse Attention', level=3)
    doc.add_paragraph(
        "Progressive refinement through three stages (f4 → f3 → f2). At each stage, creates a reverse mask "
        "(1 − σ(previous_prediction)), erasing already-confident regions and forcing the network to focus on "
        "uncertain boundary areas. Each stage produces a side output for deep supervision."
    )

    doc.add_heading('5.4 Loss Function', level=2)

    doc.add_paragraph(
        "StructureLoss (from PraNet) is used with deep supervision across main and side outputs:"
    )

    loss_p = doc.add_paragraph()
    loss_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = loss_p.add_run(
        "Loss = StructureLoss(main_pred, mask)\n"
        "     + 0.5 × StructureLoss(side1, mask)\n"
        "     + 0.3 × StructureLoss(side2, mask)\n"
        "     + 0.2 × StructureLoss(side3, mask)"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("StructureLoss combines:")
    struct_items = [
        "Weighted Binary Cross-Entropy: Edge-weighted with w = 1 + 5 × |AvgPool(mask) − mask|",
        "Weighted IoU Loss: Intersection-over-union with the same edge weighting",
    ]
    for item in struct_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "The edge weighting emphasizes polyp boundary regions, which are critical for accurate segmentation."
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 6: FAW
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('6. Fast Adaptation Weights (FAW) — Novel Contribution', level=1)

    doc.add_heading('6.1 Motivation', level=2)

    doc.add_paragraph(
        "The FAW module is the primary thesis contribution. Standard MAML adapts all model parameters "
        "(~25M) during the inner loop, which is:"
    )

    faw_issues = [
        "Computationally expensive (requires second-order gradients over millions of parameters)",
        "Memory-intensive (storing computation graphs for backpropagation-through-backpropagation)",
        "Prone to overfitting on small support sets (16 images)",
    ]
    for item in faw_issues:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "FAW solves this by concentrating adaptation into a lightweight modulation module (~100K parameters), "
        "reducing the inner-loop parameter space by approximately 250×."
    )

    doc.add_heading('6.2 Architecture', level=2)

    faw_arch_p = doc.add_paragraph()
    faw_arch_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = faw_arch_p.add_run(
        "Encoder Features [f1, f2, f3, f4]\n"
        "            ↓\n"
        "Global Average Pooling (per feature level)\n"
        "            ↓\n"
        "Concatenate → Domain Descriptor Vector\n"
        "  [256+512+1024+2048 = 3840 dims for Res2Net-50]\n"
        "            ↓\n"
        "Lightweight MLP: Linear(3840, 64) → ReLU\n"
        "            ↓\n"
        "Per-Layer Heads (3 decoder layers):\n"
        "  γ₁ = Linear(64, 32)    β₁ = Linear(64, 32)\n"
        "  γ₂ = Linear(64, 32)    β₂ = Linear(64, 32)\n"
        "  γ₃ = Linear(64, 32)    β₃ = Linear(64, 32)\n"
        "            ↓\n"
        "FiLM Modulation: feature' = γ × feature + β"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('6.3 Key Design Decisions', level=2)

    add_styled_table(doc,
        ["Decision", "Rationale"],
        [
            ["FiLM modulation", "Channel-wise affine transformation is expressive yet lightweight; proven effective for conditioning in visual reasoning"],
            ["Global Average Pooling", "Extracts domain-level statistics (mean activation per channel) capturing imaging characteristics like brightness, contrast, colour distribution"],
            ["Identity initialisation", "γ=1, β=0 at start ensures FAW has no effect initially; the model begins as a standard segmentation network and gradually learns to modulate"],
            ["MLP with hidden_dim=64", "Bottleneck design keeps parameters minimal while providing sufficient capacity to capture inter-domain differences"],
            ["Modulation at decoder level", "Decoder features are semantically richer and more task-specific; modulating here is more efficient than at the encoder"],
        ]
    )

    doc.add_heading('6.4 Why FAW Works', level=2)

    doc.add_paragraph("During meta-learning:")

    faw_why = [
        "The encoder learns to extract domain-invariant features across all training centres",
        "The FAW module learns to generate domain-specific modulations that adjust decoder features",
        "The decoder learns to produce accurate segmentations given properly modulated features",
    ]
    for i, item in enumerate(faw_why, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(item)

    doc.add_paragraph(
        "When encountering a new centre, only FAW parameters need to adapt, which captures "
        "the new domain's imaging characteristics via global statistics, generates appropriate "
        "modulations, and achieves this in just 5 gradient steps on 16 support images."
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 7: META-LEARNING FRAMEWORK
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('7. Meta-Learning Framework', level=1)

    doc.add_heading('7.1 MAML (Model-Agnostic Meta-Learning)', level=2)

    doc.add_paragraph(
        "GMLFNet uses MAML for episodic meta-training, implemented via the learn2learn library."
    )

    doc.add_heading('Training Protocol', level=3)

    protocol_steps = [
        "Sample one episode (one task per training centre, 3 tasks total)",
        "INNER LOOP — For each task: Clone model → Freeze encoder/decoder (only FAW trainable) → Adapt FAW for K=5 gradient steps on 16 support images → Unfreeze all",
        "OUTER LOOP — Forward pass on 16 query images using adapted model → Compute query loss",
        "Average query losses across all tasks",
        "Backpropagate through the entire process and update ALL parameters",
    ]
    for i, step in enumerate(protocol_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i}: ")
        run.bold = True
        p.add_run(step)

    doc.add_heading('Key Hyperparameters', level=3)

    add_styled_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["Inner learning rate", "0.01", "Step size for FAW adaptation in inner loop"],
            ["Inner steps (K)", "5", "Number of gradient steps in inner loop"],
            ["Outer learning rate", "0.001", "Step size for meta-parameter update"],
            ["Support size", "16", "Images per task for inner-loop adaptation"],
            ["Query size", "16", "Images per task for outer-loop evaluation"],
            ["Tasks per batch", "3", "One task per training centre"],
            ["First-order (FOMAML)", "True", "First-order approximation by default"],
        ]
    )

    doc.add_heading('7.2 Selective Adaptation Strategy', level=2)

    doc.add_paragraph("The key innovation in GMLFNet's meta-learning approach:")

    add_styled_table(doc,
        ["Component", "Inner Loop", "Outer Loop"],
        [
            ["Encoder (~25M params)", "Frozen", "Updated"],
            ["FAW (~100K params)", "Adapted", "Updated"],
            ["Decoder (~500K params)", "Frozen", "Updated"],
        ]
    )

    doc.add_paragraph("")
    selective_benefits = [
        "Reduces inner-loop computation by ~250× (100K vs 25M parameters)",
        "Prevents overfitting to small support sets",
        "Enables feasible training on consumer GPUs (16GB VRAM)",
        "Allows use of FOMAML without significant quality loss",
    ]
    for item in selective_benefits:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.3 FOMAML vs Full Second-Order MAML', level=2)

    add_styled_table(doc,
        ["Aspect", "FOMAML", "Second-Order MAML"],
        [
            ["Memory", "Low (no Hessian)", "High (stores computation graph)"],
            ["Speed", "Fast", "2–3× slower"],
            ["Quality", "Slightly lower", "Marginally better"],
            ["Default", "Yes (recommended)", "Available for powerful GPUs"],
        ]
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Since FAW-only adaptation already limits the inner-loop parameter space, the gap between "
        "FOMAML and full second-order MAML is minimal in practice."
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 8: ADVANTAGES
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('8. Advantages Over Existing Models', level=1)

    doc.add_heading('8.1 Comparison with Standard Segmentation Models', level=2)

    add_styled_table(doc,
        ["Aspect", "Standard Models (U-Net, PraNet, etc.)", "GMLFNet"],
        [
            ["Training paradigm", "Supervised on pooled data", "Meta-learning with episodic training"],
            ["Domain adaptation", "None (or fine-tuning)", "Built-in via MAML + FAW"],
            ["New centre deployment", "Requires retraining or fine-tuning", "Adapts in 5 gradient steps on 16 images"],
            ["Cross-centre performance", "Degrades significantly", "Robust generalization"],
            ["Adaptation speed", "Hours of fine-tuning", "Seconds of adaptation"],
        ]
    )

    doc.add_heading('8.2 Comparison with Existing Meta-Learning Approaches', level=2)

    add_styled_table(doc,
        ["Aspect", "Standard MAML", "GMLFNet (MAML + FAW)"],
        [
            ["Inner-loop parameters", "All (~25M)", "FAW only (~100K)"],
            ["Inner-loop memory", "Very high", "Low"],
            ["Overfitting risk", "High (many params, few samples)", "Low (constrained adaptation)"],
            ["Adaptation quality", "Good but unstable", "Stable and targeted"],
            ["GPU requirement", "32GB+ for second-order", "16GB sufficient"],
        ]
    )

    doc.add_heading('8.3 Comparison with Domain Adaptation Methods', level=2)

    add_styled_table(doc,
        ["Aspect", "Domain Adaptation (DANN, CycleGAN)", "GMLFNet"],
        [
            ["Target domain access", "Required during training", "Not required (zero-shot capable)"],
            ["New domain handling", "Must retrain", "Adapts with few samples"],
            ["Number of domains", "Usually 2 (source + target)", "Multiple simultaneously"],
            ["Architecture changes", "Domain discriminator needed", "Integrated FAW module"],
            ["Computational overhead", "High (adversarial training)", "Low (lightweight FAW)"],
        ]
    )

    doc.add_heading('8.4 Key Advantages Summary', level=2)

    advantages = [
        "Rapid adaptation: Adapts to new centres in seconds with 5 gradient steps",
        "Data efficient: Only needs 16 support images from the target centre",
        "Memory efficient: FAW-only adaptation reduces GPU memory by ~250× vs full MAML",
        "Zero-shot capable: Competitive performance even without adaptation",
        "Architecture agnostic: FAW can be integrated with any encoder-decoder architecture",
        "Dual backbone support: Works with both CNN (Res2Net-50) and Transformer (PVTv2-B2) encoders",
        "Clinical relevance: Addresses the real-world problem of deploying AI across hospital networks",
        "Reproducible: Full codebase with Kaggle notebook for GPU training",
    ]
    for i, item in enumerate(advantages, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(item)

    # ════════════════════════════════════════════════════════════════
    # SECTION 9: DATASETS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('9. Datasets Used', level=1)

    doc.add_heading('9.1 Overview', level=2)

    doc.add_paragraph(
        "Five standard polyp segmentation benchmark datasets are used, representing images from "
        "different medical centres across Europe:"
    )

    add_styled_table(doc,
        ["Dataset", "Images", "Resolution", "Source Institution", "Country"],
        [
            ["Kvasir-SEG", "1,000", "332×487 to 1920×1072", "Vestre Viken Health Trust", "Norway"],
            ["CVC-ClinicDB", "612", "384×288", "Hospital Clinic Barcelona", "Spain"],
            ["CVC-ColonDB", "380", "574×500", "CVC Barcelona", "Spain"],
            ["ETIS-LaribPolypDB", "196", "1225×966", "LARIB, Clermont-Ferrand", "France"],
            ["CVC-300", "60", "574×500", "CVC Barcelona", "Spain"],
            ["Total", "2,248", "", "", ""],
        ]
    )

    doc.add_heading('9.2 Data Split Protocol', level=2)

    doc.add_paragraph(
        "Following the standard multi-centre polyp segmentation evaluation protocol:"
    )

    add_styled_table(doc,
        ["Set", "Centres", "Total Images", "Purpose"],
        [
            ["Training", "Kvasir, CVC-ClinicDB, CVC-ColonDB", "1,992", "Meta-learning training (episodic)"],
            ["Testing", "ETIS-LaribPolypDB, CVC-300", "256", "Zero-shot and few-shot evaluation"],
        ]
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "This split ensures that test centres are completely unseen during training, providing a fair "
        "evaluation of cross-centre generalization."
    )

    doc.add_heading('9.3 Dataset Characteristics and Domain Shift', level=2)

    doc.add_paragraph("Each dataset exhibits distinct visual characteristics:")

    ds_chars = [
        ("Kvasir-SEG", "High variability in polyp size and appearance; green-tinted images; diverse polyp morphology"),
        ("CVC-ClinicDB", "Consistent imaging quality; sequence frames from video colonoscopy; relatively uniform lighting"),
        ("CVC-ColonDB", "Mixed quality; includes challenging cases with flat polyps; some motion blur"),
        ("ETIS-LaribPolypDB", "High resolution; significant variation in polyp texture; different endoscope model than training sets"),
        ("CVC-300", "Very small dataset; low-resolution images; challenging boundary cases"),
    ]
    for name, desc in ds_chars:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading('9.4 Data Preprocessing and Augmentation', level=2)

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run("Preprocessing:")
    run.bold = True

    preproc_items = [
        "All images resized to 352 × 352 pixels",
        "Masks binarized at threshold 128",
        "ImageNet normalization: mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]",
    ]
    for item in preproc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run("Training Augmentations:")
    run.bold = True

    aug_items = [
        "Horizontal flip (p=0.5)",
        "Vertical flip (p=0.5)",
        "Random 90-degree rotation (p=0.5)",
        "Colour jitter: brightness=0.2, contrast=0.2, saturation=0.2, hue=0.1 (p=0.5)",
        "Gaussian blur: kernel 3–7 (p=0.3)",
    ]
    for item in aug_items:
        doc.add_paragraph(item, style='List Bullet')

    # ════════════════════════════════════════════════════════════════
    # SECTION 10: TRAINING PIPELINE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('10. Training Pipeline', level=1)

    doc.add_heading('10.1 Meta-Learning Training Schedule', level=2)

    training_items = [
        "Epochs 1–10: Warmup phase (linear learning rate warmup to outer_lr = 0.001)",
        "Epochs 11–200: Cosine annealing learning rate schedule (outer_lr → 1×10⁻⁶)",
        "Each epoch: Sample N episodes, each with 3 tasks (one per training centre)",
        "Inner loop: 5 steps, FAW-only adaptation, learning rate = 0.01",
        "Outer loop: Adam optimiser with gradient clipping (max norm = 1.0)",
        "Evaluation every 10 epochs on test centres",
        "Best model saved by mean Dice score across test centres",
    ]
    for item in training_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('10.2 Infrastructure', level=2)

    add_styled_table(doc,
        ["Environment", "Use Case", "Hardware"],
        [
            ["Kaggle Notebooks", "Full training and experiments", "T4 GPU (16GB VRAM), 30h/week"],
            ["Google Colab", "Alternative training environment", "T4 GPU (16GB VRAM), ~12h sessions"],
            ["Local PC", "Development and debugging", "CPU only, 64GB RAM"],
        ]
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 11: EVALUATION METHODOLOGY
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('11. Evaluation Methodology', level=1)

    doc.add_heading('11.1 Evaluation Modes', level=2)

    doc.add_heading('Zero-Shot Generalization', level=3)
    doc.add_paragraph(
        "Direct inference on unseen test centres without any adaptation. Tests the model's "
        "inherent ability to generalize across domains. This is the primary measure of domain robustness."
    )

    doc.add_heading('Few-Shot Adaptation', level=3)
    doc.add_paragraph(
        "K support images from the target centre are used for inner-loop adaptation (K=5 steps). "
        "Remaining images are used for evaluation. Tests the model's ability to rapidly specialise "
        "to a new domain and demonstrates the value of the FAW module."
    )

    doc.add_heading('11.2 Metrics', level=2)

    doc.add_paragraph("Eight standard metrics used in polyp segmentation benchmarking:")

    add_styled_table(doc,
        ["Metric", "Range", "Description"],
        [
            ["Dice Coefficient", "[0, 1]", "Overlap between prediction and ground truth (primary metric)"],
            ["IoU (Jaccard)", "[0, 1]", "Intersection over union"],
            ["Precision", "[0, 1]", "Fraction of predicted polyp pixels that are correct"],
            ["Recall (Sensitivity)", "[0, 1]", "Fraction of actual polyp pixels that are detected"],
            ["F-measure", "[0, 1]", "Weighted harmonic mean of precision and recall (β=0.3)"],
            ["MAE", "[0, 1]", "Mean absolute error (lower is better)"],
            ["S-measure", "[0, 1]", "Structural similarity (object + region awareness)"],
            ["E-measure", "[0, 1]", "Enhanced alignment measure (local + global accuracy)"],
        ]
    )

    doc.add_heading('11.3 Expected Performance Targets', level=2)

    add_styled_table(doc,
        ["Centre", "Type", "Target Dice (Zero-Shot)", "Target Dice (Few-Shot)"],
        [
            ["Kvasir", "Seen", "> 0.85", "> 0.90"],
            ["CVC-ClinicDB", "Seen", "> 0.80", "> 0.88"],
            ["CVC-ColonDB", "Seen", "> 0.75", "> 0.82"],
            ["ETIS-LaribPolypDB", "Unseen", "> 0.65", "> 0.75"],
            ["CVC-300", "Unseen", "> 0.70", "> 0.80"],
        ]
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 12: EXPERIMENTAL DESIGN
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('12. Experimental Design', level=1)

    doc.add_heading('12.1 Main Experiments', level=2)

    add_styled_table(doc,
        ["ID", "Method", "Description"],
        [
            ["E1", "Baseline", "Standard supervised training, no meta-learning, no FAW"],
            ["E2", "MAML (full)", "MAML with all parameters adapted in inner loop"],
            ["E3", "GMLFNet (all params)", "MAML + FAW, adapt all parameters in inner loop"],
            ["E4", "GMLFNet (FAW-only) ★", "MAML + FAW, adapt only FAW in inner loop (proposed)"],
            ["E5", "FOMAML + FAW", "First-order MAML with FAW-only adaptation"],
        ]
    )

    doc.add_heading('12.2 Ablation Studies', level=2)

    add_styled_table(doc,
        ["Ablation", "Variants", "Purpose"],
        [
            ["Inner-loop steps", "K = {1, 3, 5, 10}", "Optimal adaptation depth"],
            ["FAW hidden dimension", "{32, 64, 128}", "Capacity vs. efficiency trade-off"],
            ["Support set size", "{4, 8, 16, 32}", "Data efficiency of adaptation"],
            ["Backbone architecture", "Res2Net-50 vs. PVTv2-B2", "CNN vs. Transformer comparison"],
            ["FAW ablation", "With FAW vs. without FAW", "Validates FAW contribution"],
            ["Leave-one-centre-out", "5 runs (hold out each centre)", "Robustness across configurations"],
        ]
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 13: TECHNICAL SPECIFICATIONS
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('13. Technical Specifications', level=1)

    doc.add_heading('13.1 Software Dependencies', level=2)

    add_styled_table(doc,
        ["Package", "Version", "Purpose"],
        [
            ["PyTorch", ">= 2.0.0", "Deep learning framework"],
            ["torchvision", ">= 0.15.0", "Image transforms and models"],
            ["learn2learn", ">= 0.2.0", "MAML implementation"],
            ["timm", ">= 0.9.0", "Pretrained backbones (Res2Net, PVTv2)"],
            ["albumentations", ">= 1.3.0", "Image augmentation pipeline"],
            ["OpenCV", ">= 4.7.0", "Image I/O and processing"],
            ["NumPy", ">= 1.24.0", "Numerical computing"],
            ["PyYAML", ">= 6.0", "Configuration loading"],
            ["TensorBoard", ">= 2.13.0", "Training visualisation"],
            ["matplotlib", ">= 3.7.0", "Result plotting"],
            ["gdown", ">= 4.7.0", "Google Drive dataset download"],
        ]
    )

    doc.add_heading('13.2 Model Size', level=2)

    add_styled_table(doc,
        ["Component", "Res2Net-50", "PVTv2-B2"],
        [
            ["Encoder", "~25,000,000", "~25,000,000"],
            ["FAW", "~100,000", "~60,000"],
            ["Decoder", "~500,000", "~200,000"],
            ["Total", "~25,600,000", "~25,260,000"],
        ]
    )

    doc.add_heading('13.3 Computational Requirements', level=2)

    add_styled_table(doc,
        ["Requirement", "Specification"],
        [
            ["GPU VRAM", "16GB minimum (T4 or better)"],
            ["Training time", "~3–6 hours for 200 epochs on T4"],
            ["Inference time", "~30ms per image on T4 GPU"],
            ["Adaptation time", "~2 seconds (5 steps on 16 images)"],
            ["Disk space", "~2GB for datasets, ~100MB per checkpoint"],
        ]
    )

    # ════════════════════════════════════════════════════════════════
    # SECTION 14: PROJECT STRUCTURE
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('14. Project Structure', level=1)

    structure_p = doc.add_paragraph()
    run = structure_p.add_run(
        "GMLFNet/\n"
        "├── README.md                       Quick-start guide\n"
        "├── requirements.txt                Python dependencies\n"
        "├── .gitignore                      Git exclusions\n"
        "├── docs/\n"
        "│   └── PROJECT_DOCUMENT.md         This document (Markdown version)\n"
        "├── configs/\n"
        "│   └── default.yaml                Training configuration\n"
        "├── data/\n"
        "│   ├── datasets.py                 PolypCenterDataset class\n"
        "│   ├── augmentations.py            Train/test augmentation pipelines\n"
        "│   ├── meta_sampler.py             Episodic task sampler for MAML\n"
        "│   └── download.py                 Dataset download script\n"
        "├── models/\n"
        "│   ├── backbone.py                 Res2Net-50 and PVTv2-B2 encoders\n"
        "│   ├── decoder.py                  Multi-scale decoder (RFB + RA)\n"
        "│   ├── gmlf_net.py                 Full GMLFNet architecture\n"
        "│   ├── fast_adapt_weights.py       FAW module (thesis contribution)\n"
        "│   └── losses.py                   StructureLoss + deep supervision\n"
        "├── trainers/\n"
        "│   ├── meta_trainer.py             MAML meta-learning trainer\n"
        "│   ├── baseline_trainer.py         Standard supervised trainer\n"
        "│   └── evaluator.py               Zero-shot and few-shot evaluation\n"
        "├── utils/\n"
        "│   ├── metrics.py                  8 segmentation metrics\n"
        "│   ├── visualization.py            Result plotting utilities\n"
        "│   ├── logging_utils.py            TensorBoard/W&B logging\n"
        "│   └── misc.py                     Config, seeding, checkpointing\n"
        "├── scripts/\n"
        "│   ├── train_meta.py               Entry: meta-learning training\n"
        "│   ├── train_baseline.py           Entry: baseline training\n"
        "│   ├── evaluate.py                 Entry: model evaluation\n"
        "│   └── ablation.py                 Entry: ablation studies\n"
        "└── notebooks/\n"
        "    ├── kaggle_train.ipynb           Self-contained Kaggle notebook\n"
        "    ├── data_exploration.ipynb       Dataset visualisation\n"
        "    └── results_analysis.ipynb       Post-training analysis"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # ════════════════════════════════════════════════════════════════
    # SECTION 15: REFERENCES
    # ════════════════════════════════════════════════════════════════
    doc.add_heading('15. References', level=1)

    doc.add_heading('Foundational Works', level=2)

    refs_foundational = [
        "[1] Finn, C., Abbeel, P., & Levine, S. (2017). Model-Agnostic Meta-Learning for Fast Adaptation of Deep Networks. ICML.",
        "[2] Perez, E., Strub, F., de Vries, H., Dumoulin, V., & Courville, A. (2018). FiLM: Visual Reasoning with a General Conditioning Layer. AAAI.",
        "[3] Fan, D.P., Ji, G.P., Zhou, T., Chen, G., Fu, H., Shen, J., & Shao, L. (2020). PraNet: Parallel Reverse Attention Network for Polyp Segmentation. MICCAI.",
    ]
    for ref in refs_foundational:
        doc.add_paragraph(ref)

    doc.add_heading('Backbone Architectures', level=2)

    refs_backbone = [
        "[4] Gao, S.H., Cheng, M.M., Zhao, K., Zhang, X.Y., Yang, M.H., & Torr, P. (2019). Res2Net: A New Multi-scale Backbone Architecture. IEEE TPAMI.",
        "[5] Wang, W., Xie, E., Li, X., Fan, D.P., Song, K., Liang, D., Lu, T., Luo, P., & Shao, L. (2022). PVT v2: Improved Baselines with Pyramid Vision Transformer. Computational Visual Media.",
    ]
    for ref in refs_backbone:
        doc.add_paragraph(ref)

    doc.add_heading('Polyp Segmentation', level=2)

    refs_polyp = [
        "[6] Dong, B., Wang, W., Fan, D.P., Li, J., Fu, H., & Shao, L. (2021). Polyp-PVT: Polyp Segmentation with Pyramid Vision Transformers. arXiv.",
    ]
    for ref in refs_polyp:
        doc.add_paragraph(ref)

    doc.add_heading('Evaluation Metrics', level=2)

    refs_metrics = [
        "[7] Fan, D.P., Cheng, M.M., Liu, Y., Li, T., & Borji, A. (2017). Structure-measure: A New Way to Evaluate Foreground Maps. ICCV.",
        "[8] Fan, D.P., Gong, C., Cao, Y., Ren, B., Cheng, M.M., & Borji, A. (2018). Enhanced-alignment Measure for Binary Foreground Map Evaluation. IJCAI.",
    ]
    for ref in refs_metrics:
        doc.add_paragraph(ref)

    doc.add_heading('Datasets', level=2)

    refs_datasets = [
        "[9] Jha, D., Smedsrud, P.H., Riegler, M.A., et al. (2020). Kvasir-SEG: A Segmented Polyp Dataset. MMM.",
        "[10] Bernal, J., Sanchez, F.J., et al. (2015). WM-DOVA Maps for Accurate Polyp Highlighting in Colonoscopy. TMI.",
        "[11] Tajbakhsh, N., Gurudu, S.R., & Liang, J. (2015). Automated Polyp Detection in Colonoscopy Videos Using Shape and Context Information. TMI.",
        "[12] Silva, J., Histace, A., et al. (2014). Toward Embedded Detection of Polyps in WCE Images for Early Diagnosis of Colorectal Cancer. IJCARS.",
    ]
    for ref in refs_datasets:
        doc.add_paragraph(ref)

    # ── Footer ──
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_text.add_run(
        "This document is part of the GMLFNet research project for a Master's thesis in Computer Science.\n"
        "Repository: https://github.com/bsaccoh/GMLFNet-Research-Project"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True

    # ── Save ──
    output_path = Path(__file__).parent / "GMLFNet_Project_Document.docx"
    doc.save(str(output_path))
    print(f"Document saved to: {output_path}")
    print(f"Size: {output_path.stat().st_size / 1024:.1f} KB")
    return output_path


if __name__ == "__main__":
    create_document()
